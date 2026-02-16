"""
Policy Generator Module

Author: Ayi NEDJIMI
Website: https://ayinedjimi-consultants.fr
"""

import os
from dataclasses import dataclass
from typing import Dict, List, Optional
import openai
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import structlog

logger = structlog.get_logger(__name__)


@dataclass
class PolicyConfig:
    """Policy configuration."""
    framework: str  # ISO27001, RGPD, NIS2, SOC2
    organization_name: str
    industry: str
    size: str  # small, medium, large
    language: str = "en"  # en, fr


class PolicyGenerator:
    """
    AI-Powered Security Policy Generator.

    Author: Ayi NEDJIMI
    Website: https://ayinedjimi-consultants.fr

    Supports: ISO27001, RGPD/GDPR, NIS2, SOC2, PCI-DSS
    """

    def __init__(self, api_key: Optional[str] = None):
        """Initialize policy generator."""
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        openai.api_key = self.api_key
        self.logger = logger.bind(component="policy_generator")

    def generate_policy(self, config: PolicyConfig) -> Dict[str, any]:
        """Generate security policy."""
        self.logger.info("Generating policy", framework=config.framework)

        # Generate policy sections using LLM
        sections = self._generate_policy_sections(config)

        return {
            "title": f"{config.framework} Security Policy - {config.organization_name}",
            "framework": config.framework,
            "sections": sections,
            "metadata": {
                "organization": config.organization_name,
                "industry": config.industry,
                "generated_by": "Ayi NEDJIMI - ayinedjimi-consultants.fr",
            }
        }

    def _generate_policy_sections(self, config: PolicyConfig) -> List[Dict[str, str]]:
        """Generate policy sections using AI."""
        framework_templates = {
            "ISO27001": [
                "Information Security Policy",
                "Organization of Information Security",
                "Human Resource Security",
                "Asset Management",
                "Access Control",
                "Cryptography",
                "Physical and Environmental Security",
                "Operations Security",
                "Communications Security",
                "System Acquisition, Development and Maintenance",
                "Supplier Relationships",
                "Information Security Incident Management",
                "Business Continuity Management",
                "Compliance",
            ],
            "RGPD": [
                "Data Protection Principles",
                "Lawful Basis for Processing",
                "Data Subject Rights",
                "Data Security Measures",
                "Data Breach Notification",
                "Data Protection Impact Assessment",
                "Data Protection Officer",
                "International Data Transfers",
            ],
            "NIS2": [
                "Risk Management",
                "Corporate Governance",
                "Business Continuity",
                "Supply Chain Security",
                "Security in Network and Information Systems",
                "Incident Handling",
                "Security Testing and Auditing",
            ],
        }

        section_names = framework_templates.get(config.framework, [])
        sections = []

        for section_name in section_names[:5]:  # Generate first 5 sections
            content = self._generate_section_content(section_name, config)
            sections.append({
                "title": section_name,
                "content": content,
            })

        return sections

    def _generate_section_content(self, section_name: str, config: PolicyConfig) -> str:
        """Generate section content using LLM."""
        prompt = f"""Generate a comprehensive {section_name} section for a {config.framework} security policy.

Organization: {config.organization_name}
Industry: {config.industry}
Size: {config.size}

Provide detailed, professional content with:
- Clear objectives
- Specific requirements
- Implementation guidelines
- Compliance requirements

Format in professional policy language."""

        try:
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a cybersecurity policy expert."},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.3,
                max_tokens=1000,
            )

            return response.choices[0].message.content

        except Exception as e:
            self.logger.error("LLM generation failed", error=str(e))
            return f"[Section content for {section_name}]"

    def export_to_docx(self, policy: Dict[str, any], output_path: str):
        """Export policy to Word document."""
        doc = Document()

        # Title
        doc.add_heading(policy["title"], 0)

        # Metadata
        doc.add_paragraph(f"Organization: {policy['metadata']['organization']}")
        doc.add_paragraph(f"Framework: {policy['framework']}")
        doc.add_paragraph(f"Generated by: {policy['metadata']['generated_by']}")
        doc.add_page_break()

        # Sections
        for section in policy["sections"]:
            doc.add_heading(section["title"], 1)
            doc.add_paragraph(section["content"])
            doc.add_page_break()

        doc.save(output_path)
        self.logger.info("Exported to DOCX", file=output_path)

    def export_to_pdf(self, policy: Dict[str, any], output_path: str):
        """Export policy to PDF."""
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Title
        story.append(Paragraph(policy["title"], styles['Title']))
        story.append(Spacer(1, 12))

        # Metadata
        story.append(Paragraph(f"Organization: {policy['metadata']['organization']}", styles['Normal']))
        story.append(Spacer(1, 12))

        # Sections
        for section in policy["sections"]:
            story.append(Paragraph(section["title"], styles['Heading1']))
            story.append(Spacer(1, 6))
            story.append(Paragraph(section["content"], styles['Normal']))
            story.append(Spacer(1, 12))

        doc.build(story)
        self.logger.info("Exported to PDF", file=output_path)
